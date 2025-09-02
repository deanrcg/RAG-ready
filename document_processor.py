#!/usr/bin/env python3
"""
Multi-format document processor for RAG applications.
Supports PDF, DOCX, TXT, MD, and Excel files.
"""

import os
import re
from typing import Dict, List, Tuple, Optional
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

class DocumentProcessor:
    """Handles multiple document formats and extracts text content."""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._extract_txt,
            '.md': self._extract_md,
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_excel,
            '.xls': self._extract_excel,
        }
    
    def can_process(self, file_path: str) -> bool:
        """Check if the file format is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_formats
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text content from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (text_content, metadata)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")
        
        return self.supported_formats[ext](file_path)
    
    def _extract_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content, {"format": "txt", "encoding": "utf-8"}
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
            return content, {"format": "txt", "encoding": "latin-1"}
    
    def _extract_md(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from Markdown files."""
        content, metadata = self._extract_txt(file_path)
        metadata["format"] = "markdown"
        return content, metadata
    
    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from PDF files."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
        
        text_parts = []
        metadata = {"format": "pdf"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                
                # Extract PDF metadata
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    if pdf_meta.get('/Title'):
                        metadata["title"] = pdf_meta['/Title']
                    if pdf_meta.get('/Author'):
                        metadata["author"] = pdf_meta['/Author']
                    if pdf_meta.get('/Subject'):
                        metadata["subject"] = pdf_meta['/Subject']
        
        except Exception as e:
            raise ValueError(f"Error processing PDF {file_path}: {str(e)}")
        
        return "\n\n".join(text_parts), metadata
    
    def _extract_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            metadata = {"format": "docx"}
            
            # Extract document properties
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.subject:
                metadata["subject"] = doc.core_properties.subject
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts), metadata
        
        except Exception as e:
            raise ValueError(f"Error processing DOCX {file_path}: {str(e)}")
    
    def _extract_excel(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from Excel files."""
        if not EXCEL_AVAILABLE:
            raise ImportError("pandas and openpyxl are required for Excel processing. Install with: pip install pandas openpyxl")
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            metadata = {"format": "excel", "sheets": excel_file.sheet_names}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                if not df.empty:
                    # Convert DataFrame to text
                    sheet_text = f"--- Sheet: {sheet_name} ---\n"
                    sheet_text += df.to_string(index=False)
                    text_parts.append(sheet_text)
            
            return "\n\n".join(text_parts), metadata
        
        except Exception as e:
            raise ValueError(f"Error processing Excel {file_path}: {str(e)}")
    
    def batch_process(self, folder_path: str) -> List[Tuple[str, str, Dict]]:
        """
        Process all supported files in a folder.
        
        Args:
            folder_path: Path to the folder containing documents
            
        Returns:
            List of tuples: (file_path, text_content, metadata)
        """
        results = []
        folder = Path(folder_path)
        
        if not folder.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")
        
        for file_path in folder.rglob("*"):
            if file_path.is_file() and self.can_process(str(file_path)):
                try:
                    text, metadata = self.extract_text(str(file_path))
                    results.append((str(file_path), text, metadata))
                except Exception as e:
                    print(f"Warning: Could not process {file_path}: {e}")
        
        return results

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove page markers
    text = re.sub(r'--- Page \d+ ---', '', text)
    text = re.sub(r'--- Sheet: .*? ---', '', text)
    
    # Clean up bullet points
    text = re.sub(r'^\s*[•\-\*]\s*', '• ', text, flags=re.MULTILINE)
    
    # Remove empty lines
    text = re.sub(r'\n\s*\n', '\n', text)
    
    return text.strip()

if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    
    # Test with a file
    test_file = "content/intro-to-hse/01-overview.md"
    if os.path.exists(test_file):
        text, metadata = processor.extract_text(test_file)
        print(f"Extracted text from {test_file}:")
        print(f"Metadata: {metadata}")
        print(f"Text length: {len(text)} characters")
        print(f"First 200 chars: {text[:200]}...")
